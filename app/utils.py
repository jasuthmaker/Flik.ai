import os
import pdfplumber
import pytesseract
from PIL import Image
from werkzeug.utils import secure_filename
from app.ai_processor import AIProcessor

# Optional Google integrations (lazy import pattern)
try:
    from google.cloud import documentai
    _has_docai = True
except Exception:
    _has_docai = False

def extract_text_from_pdf(file_path):
    """Extract text from PDF file using pdfplumber (local)."""
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_image(file_path):
    """Extract text from image using OCR (pytesseract)."""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from image: {e}")
        return ""

def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        print(f"Error reading text file: {e}")
        return ""

def extract_text_from_docx(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_with_document_ai(file_path):
    """Optional: Use Google Document AI if configured; otherwise return ''."""
    if not _has_docai:
        return ""
    try:
        project_id = os.getenv("GOOGLE_PROJECT_ID")
        location = os.getenv("GOOGLE_LOCATION", "us")
        processor_id = os.getenv("GOOGLE_PROCESSOR_ID")
        if not (project_id and processor_id):
            return ""
        client = documentai.DocumentProcessorServiceClient()
        name = client.processor_path(project_id, location, processor_id)
        with open(file_path, "rb") as f:
            raw_document = documentai.RawDocument(content=f.read(), mime_type="application/pdf")
        request = documentai.ProcessRequest(name=name, raw_document=raw_document)
        result = client.process_document(request=request)
        return (result.document.text or '').strip()
    except Exception as e:
        print(f"Document AI extraction failed: {e}")
        return ""

def extract_text_from_file(file_path, file_type):
    file_type = file_type.lower()

    # Prefer GCP Document AI for PDFs if configured
    if file_type == 'pdf':
        cloud_text = extract_text_with_document_ai(file_path)
        if cloud_text:
            return cloud_text
        return extract_text_from_pdf(file_path)
    elif file_type in ['png', 'jpg', 'jpeg']:
        return extract_text_from_image(file_path)
    elif file_type == 'txt':
        return extract_text_from_txt(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    else:
        return ""

def get_file_size(file_path):
    return os.path.getsize(file_path)

def format_file_size(size_bytes):
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    return f"{size_bytes:.1f} {size_names[i]}"

def process_document_with_ai(text, filename=""):
    try:
        ai_processor = AIProcessor()
        category, appointments_todos = ai_processor.process_document(text, filename)
        return category, appointments_todos
    except Exception as e:
        print(f"Error in AI processing: {e}")
        return "Other", []
